from __future__ import annotations

import copy
import io
import json
import re
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =========================
# Existing resume processor
# =========================

@dataclass
class ResumeLine:
    index: int
    text: str
    char_count: int
    source: str  # "paragraph" | "table"
    para_index: int
    table_ref: Optional[Tuple[int, int, int]] = None
    is_empty: bool = False


class ResumeProcessor:
    """
    Parses a DOCX resume into line-level units and supports precise replacement
    while preserving formatting as much as possible.
    """

    def __init__(self, docx_bytes: bytes):
        self._original_bytes = docx_bytes
        self.doc = Document(io.BytesIO(docx_bytes))
        self._lines: List[ResumeLine] = []
        self._line_lookup: Dict[int, Dict[str, Any]] = {}
        self._parse()

    def _parse(self) -> None:
        idx = 0

        for pi, para in enumerate(self.doc.paragraphs):
            text = self._clean_text(para.text)
            self._lines.append(
                ResumeLine(
                    index=idx,
                    text=text,
                    char_count=len(text),
                    source="paragraph",
                    para_index=pi,
                    is_empty=(text == ""),
                )
            )
            self._line_lookup[idx] = {"source": "paragraph", "para_index": pi}
            idx += 1

        for ti, table in enumerate(self.doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    text = self._clean_text(cell.text)
                    self._lines.append(
                        ResumeLine(
                            index=idx,
                            text=text,
                            char_count=len(text),
                            source="table",
                            para_index=-1,
                            table_ref=(ti, ri, ci),
                            is_empty=(text == ""),
                        )
                    )
                    self._line_lookup[idx] = {
                        "source": "table",
                        "table_ref": (ti, ri, ci),
                    }
                    idx += 1

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\xa0", " ")
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    def get_all_lines(self, include_empty: bool = True) -> List[dict]:
        lines = self._lines if include_empty else [l for l in self._lines if not l.is_empty]
        return [
            {
                "index": l.index,
                "text": l.text,
                "char_count": l.char_count,
                "source": l.source,
            }
            for l in lines
        ]

    def get_line(self, line_index: int) -> Optional[ResumeLine]:
        if 0 <= line_index < len(self._lines):
            return self._lines[line_index]
        return None

    def get_context_window(self, line_index: int, window: int = 1) -> dict:
        start = max(0, line_index - window)
        end = min(len(self._lines), line_index + window + 1)

        return {
            "target": self._lines[line_index].text if 0 <= line_index < len(self._lines) else "",
            "before": [self._lines[i].text for i in range(start, line_index)],
            "after": [self._lines[i].text for i in range(line_index + 1, end)],
        }

    def replace_line(self, line_index: int, new_text: str) -> bool:
        if not (0 <= line_index < len(self._lines)):
            return False

        new_text = self._clean_text(new_text)
        line = self._lines[line_index]
        lookup = self._line_lookup[line_index]

        if lookup["source"] == "paragraph":
            para = self.doc.paragraphs[lookup["para_index"]]
            self._replace_para_text(para, new_text)

        elif lookup["source"] == "table":
            ti, ri, ci = lookup["table_ref"]
            cell = self.doc.tables[ti].rows[ri].cells[ci]

            if cell.paragraphs:
                self._replace_para_text(cell.paragraphs[0], new_text)
            else:
                cell.text = new_text

        line.text = new_text
        line.char_count = len(new_text)
        line.is_empty = (new_text == "")
        return True

    @staticmethod
    def _replace_para_text(para, new_text: str) -> None:
        """
        Replace paragraph text while preserving the first meaningful run style.
        """
        template_run = None
        for run in para.runs:
            if run.text and run.text.strip():
                template_run = run
                break

        if template_run is None and para.runs:
            template_run = para.runs[0]

        rPr_clone = None
        if template_run is not None:
            rPr = template_run._r.find(qn("w:rPr"))
            if rPr is not None:
                rPr_clone = copy.deepcopy(rPr)

        p_elem = para._p

        for child in list(p_elem):
            if child.tag == qn("w:r"):
                p_elem.remove(child)

        new_run = OxmlElement("w:r")
        if rPr_clone is not None:
            new_run.append(rPr_clone)

        text_elem = OxmlElement("w:t")
        if new_text.startswith(" ") or new_text.endswith(" "):
            text_elem.set(qn("xml:space"), "preserve")
        text_elem.text = new_text

        new_run.append(text_elem)
        p_elem.append(new_run)

    def export(self) -> bytes:
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()


# =========================
# Production-ready ATS layer
# =========================

@dataclass
class JDAnalysis:
    core_role_function: str = ""
    primary_business_objective: str = ""
    required_skills: List[str] = field(default_factory=list)
    preferred_skills: List[str] = field(default_factory=list)
    required_tools: List[str] = field(default_factory=list)
    required_domain_terms: List[str] = field(default_factory=list)
    required_responsibilities: List[str] = field(default_factory=list)
    required_qualifications: List[str] = field(default_factory=list)
    preferred_qualifications: List[str] = field(default_factory=list)
    repeated_keywords: List[str] = field(default_factory=list)
    synonyms_and_variants: Dict[str, List[str]] = field(default_factory=dict)
    implied_screening_signals: List[str] = field(default_factory=list)
    metrics_kpi_language: List[str] = field(default_factory=list)
    ats_critical_terms: List[str] = field(default_factory=list)
    required_keywords: List[str] = field(default_factory=list)
    preferred_keywords: List[str] = field(default_factory=list)
    high_priority_missing: List[str] = field(default_factory=list)
    medium_priority_missing: List[str] = field(default_factory=list)
    low_priority_missing: List[str] = field(default_factory=list)
    recommended_keyword_targets: List[str] = field(default_factory=list)


@dataclass
class LineAssessment:
    relevance: str = "low"
    current_match_score: int = 0
    issues: List[str] = field(default_factory=list)


@dataclass
class RewriteOption:
    text: str
    keywords_added: List[str] = field(default_factory=list)
    score: int = 0
    why_it_works: str = ""


@dataclass
class LineSuggestion:
    line_index: int
    original: str
    assessment: LineAssessment
    options: List[RewriteOption] = field(default_factory=list)
    best_option_number: int = 1
    best_option_reason: str = ""
    risk_flags: List[str] = field(default_factory=list)


@dataclass
class GlobalResumeFeedback:
    rewrite_based_improvements: List[str] = field(default_factory=list)
    substantive_gaps: List[str] = field(default_factory=list)
    parser_risks: List[str] = field(default_factory=list)
    overall_estimated_match_score: int = 0
    what_prevents_true_100_percent_match: List[str] = field(default_factory=list)


@dataclass
class ATSReviewBundle:
    jd_analysis: JDAnalysis
    line_suggestions: List[LineSuggestion]
    global_resume_feedback: GlobalResumeFeedback


class ATSUtils:
    COMMON_PARSER_RISKS = [
        "text boxes and floating elements may not parse correctly",
        "graphics, icons, and image-based text are often ignored by ATS",
        "tables with complex merged cells can hurt parsing",
        "headers/footers may be skipped by some systems",
        "non-standard section names can weaken ATS extraction",
        "multi-column layouts can cause reading-order issues",
    ]

    ACTION_VERBS = {
        "analyzed", "led", "built", "developed", "designed", "created", "implemented",
        "improved", "optimized", "managed", "reduced", "increased", "streamlined",
        "coordinated", "supported", "executed", "delivered", "drove", "evaluated",
        "monitored", "produced", "launched", "automated", "engineered", "performed"
    }

    @staticmethod
    def clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\xa0", " ")
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    @staticmethod
    def normalize_token(token: str) -> str:
        token = ATSUtils.clean_text(token).lower()
        token = re.sub(r"[^a-z0-9+/&().# -]", "", token)
        return token.strip()

    @staticmethod
    def normalize_text(text: str) -> str:
        if not isinstance(text, str):
            return ""
        text = text.strip()
        text = text.replace("\u201c", '"').replace("\u201d", '"').replace("\u2019", "'")
        text = re.sub(r"\s+", " ", text)
        return text

    @staticmethod
    def first_word(text: str) -> str:
        normalized = ATSUtils.normalize_text(text)
        parts = normalized.split()
        return parts[0].lower() if parts else ""

    @staticmethod
    def should_preserve_first_word(text: str) -> bool:
        first = ATSUtils.first_word(text)
        return first in ATSUtils.ACTION_VERBS

    @staticmethod
    def dedupe_keep_order(items: List[str]) -> List[str]:
        seen = set()
        output = []
        for item in items:
            key = ATSUtils.normalize_token(item)
            if key and key not in seen:
                seen.add(key)
                output.append(ATSUtils.clean_text(item))
        return output

    @staticmethod
    def safe_parse_json(raw: str) -> Any:
        if raw is None:
            raise ValueError("Empty model output")

        cleaned = raw.strip()
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.IGNORECASE | re.DOTALL)
        cleaned = cleaned.replace("\u201c", '"').replace("\u201d", '"').replace("\u2019", "'")
        cleaned = re.sub(r",(\s*[\]}])", r"\1", cleaned)

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError as exc:
            raise ValueError(f"JSON parse failed: {exc}") from exc

    @staticmethod
    def find_keyword_hits(text: str, keywords: List[str]) -> List[str]:
        normalized_text = ATSUtils.normalize_token(text)
        hits = []
        for kw in keywords:
            nkw = ATSUtils.normalize_token(kw)
            if nkw and nkw in normalized_text:
                hits.append(kw)
        return ATSUtils.dedupe_keep_order(hits)

    @staticmethod
    def pretty_json(data: Any) -> str:
        return json.dumps(data, indent=2, ensure_ascii=False)


class ATSScorer:
    """
    Lightweight heuristic scorer. Strict on truth and layout safety.
    """

    WEIGHTS = {
        "keyword_relevance": 30,
        "truthfulness": 20,
        "tool_alignment": 15,
        "readability": 10,
        "specificity": 10,
        "parser_friendliness": 10,
        "length_control": 5,
    }

    @staticmethod
    def assess_line_relevance(line_text: str, jd: JDAnalysis) -> LineAssessment:
        text = ATSUtils.clean_text(line_text)
        if not text:
            return LineAssessment(relevance="low", current_match_score=0, issues=["empty line"])

        critical = jd.ats_critical_terms or []
        tools = jd.required_tools or []
        skills = jd.required_skills or []
        responsibilities = jd.required_responsibilities or []

        all_terms = ATSUtils.dedupe_keep_order(critical + tools + skills + responsibilities)
        hits = ATSUtils.find_keyword_hits(text, all_terms)

        score = min(100, int((len(hits) / max(1, min(len(all_terms), 10))) * 100))
        if score >= 65:
            relevance = "high"
        elif score >= 30:
            relevance = "medium"
        else:
            relevance = "low"

        issues = []
        if not hits:
            issues.append("missing obvious JD-aligned terminology")
        if len(text) < 25:
            issues.append("line may be too short to communicate scope or impact")
        if not re.search(r"\d", text):
            issues.append("no measurable detail present")
        if len(text.split()) < 4:
            issues.append("low specificity")

        return LineAssessment(
            relevance=relevance,
            current_match_score=score,
            issues=ATSUtils.dedupe_keep_order(issues),
        )

    @staticmethod
    def score_rewrite_option(
        original: str,
        option: str,
        jd: JDAnalysis,
        max_chars: Optional[int] = None,
        selected_keywords: Optional[List[str]] = None,
    ) -> int:
        text = ATSUtils.clean_text(option)
        original_text = ATSUtils.clean_text(original)
        selected_keywords = selected_keywords or []

        critical_terms = ATSUtils.dedupe_keep_order(
            jd.ats_critical_terms + jd.required_tools + jd.required_skills + jd.required_domain_terms
        )
        critical_hits = ATSUtils.find_keyword_hits(text, critical_terms)
        keyword_score = min(20, len(critical_hits) * 3)

        selected_hits = ATSUtils.find_keyword_hits(text, selected_keywords)
        selected_keyword_score = min(25, len(selected_hits) * 5)

        truthfulness_score = ATSScorer.WEIGHTS["truthfulness"]
        if len(text) > len(original_text):
            truthfulness_score -= 2
        if text.count(";") > 1:
            truthfulness_score -= 2

        tool_hits = ATSUtils.find_keyword_hits(text, jd.required_tools)
        tool_score = min(ATSScorer.WEIGHTS["tool_alignment"], len(tool_hits) * 5)

        readability_score = ATSScorer.WEIGHTS["readability"]
        if len(text.split()) < 4 or len(text.split()) > 35:
            readability_score -= 3
        if re.search(r"\b(results-driven|dynamic|synergy|go-getter|hardworking)\b", text, flags=re.I):
            readability_score -= 3

        specificity_score = ATSScorer.WEIGHTS["specificity"]
        if not re.search(r"\d", text):
            specificity_score -= 4
        if len(text.split()) < 5:
            specificity_score -= 2

        parser_score = ATSScorer.WEIGHTS["parser_friendliness"]
        if re.search(r"[|¦•►◆■✔]", text):
            parser_score -= 4

        selected_norm = {ATSUtils.normalize_token(k) for k in selected_keywords}
        non_selected_hits = [
            kw for kw in ATSUtils.find_keyword_hits(text, jd.ats_critical_terms)
            if ATSUtils.normalize_token(kw) not in selected_norm
        ]
        if non_selected_hits:
            parser_score -= min(5, len(non_selected_hits))

        length_score = ATSScorer.WEIGHTS["length_control"]
        if max_chars is not None and len(text) > max_chars:
            length_score = 0

        total = (
            keyword_score
            + selected_keyword_score
            + max(0, truthfulness_score)
            + tool_score
            + max(0, readability_score)
            + max(0, specificity_score)
            + max(0, parser_score)
            + max(0, length_score)
        )
        return max(0, min(100, int(total)))


class ResumeRewritePrompts:
    @staticmethod
    def build_master_ats_prompt(
        resume_lines: List[Dict[str, Any]],
        job_description: str,
    ) -> str:
        return f"""
You are an elite resume-job description matching engine built to maximize ATS alignment while staying 100% truthful.

Your job is to analyze the full job description with extreme care, identify every assessable hiring signal, compare it against the candidate's resume line by line, and generate optimized resume suggestions that maximize ATS match quality.

PRIMARY GOAL
Produce the strongest possible ATS-aligned suggestions while preserving truth, clarity, formatting intent, natural human resume language, and exact line-length safety.

NON-NEGOTIABLE RULES
1. Never invent experience, tools, certifications, metrics, industries, leadership, scope, or outcomes not supported by the resume.
2. Never exaggerate beyond what the original line can reasonably support.
3. Keep suggestions natural, concise, recruiter-friendly, and resume-appropriate.
4. Optimize for both ATS parsing and recruiter readability.
5. Treat every word in the job description as potentially important.
6. Keep rewrite suggestions layout-safe for tightly formatted resumes.
7. Do NOT claim "100% ATS match" unless the resume genuinely covers the JD's requirements.

FINAL OUTPUT FORMAT
Return ONLY valid JSON in this structure:

{{
  "jd_analysis": {{
    "core_role_function": "",
    "primary_business_objective": "",
    "required_skills": [],
    "preferred_skills": [],
    "required_tools": [],
    "required_domain_terms": [],
    "required_responsibilities": [],
    "required_qualifications": [],
    "preferred_qualifications": [],
    "repeated_keywords": [],
    "synonyms_and_variants": {{}},
    "implied_screening_signals": [],
    "metrics_kpi_language": [],
    "ats_critical_terms": [],
    "required_keywords": [],
    "preferred_keywords": [],
    "high_priority_missing": [],
    "medium_priority_missing": [],
    "low_priority_missing": [],
    "recommended_keyword_targets": []
  }},
  "line_suggestions": [],
  "global_resume_feedback": {{
    "rewrite_based_improvements": [],
    "substantive_gaps": [],
    "parser_risks": [],
    "overall_estimated_match_score": 0,
    "what_prevents_true_100_percent_match": []
  }}
}}

Resume lines:
{json.dumps(resume_lines, ensure_ascii=False)}

Job description:
{job_description}
""".strip()

    @staticmethod
    def build_jd_deep_parse_prompt(job_description: str) -> str:
        return f"""
Extract ATS-critical hiring signals from this job description.

Return ONLY valid JSON with:
{{
  "core_role_function": "",
  "primary_business_objective": "",
  "required_skills": [],
  "preferred_skills": [],
  "required_tools": [],
  "required_domain_terms": [],
  "required_responsibilities": [],
  "required_qualifications": [],
  "preferred_qualifications": [],
  "repeated_keywords": [],
  "synonyms_and_variants": {{
    "keyword_from_jd": ["variant1", "variant2"]
  }},
  "implied_screening_signals": [],
  "metrics_kpi_language": [],
  "ats_critical_terms": [],
  "required_keywords": [],
  "preferred_keywords": [],
  "high_priority_missing": [],
  "medium_priority_missing": [],
  "low_priority_missing": [],
  "recommended_keyword_targets": []
}}

Job description:
{job_description}
""".strip()

    @staticmethod
    def build_line_rewrite_prompt(
        line_index: int,
        original_line: str,
        job_description: str,
        selected_keywords: Optional[List[str]] = None,
        surrounding_before: Optional[List[str]] = None,
        surrounding_after: Optional[List[str]] = None,
        max_chars: Optional[int] = None,
        n_options: int = 5,
    ) -> str:
        surrounding_before = surrounding_before or []
        surrounding_after = surrounding_after or []
        selected_keywords = selected_keywords or []

        first_word = ATSUtils.first_word(original_line)
        preserve_first_word = ATSUtils.should_preserve_first_word(original_line)

        options_schema = ",\n".join([
            """    {
      "text": "",
      "keywords_added": [],
      "score": 0,
      "why_it_works": ""
    }""" for _ in range(n_options)
        ])

        return f"""
Rewrite this single resume line to maximize truthful ATS alignment.

CRITICAL RULES
1. Use ONLY keywords from the SELECTED KEYWORDS list when adding ATS language.
2. Do not add any keyword unless it fits truthfully.
3. Do not invent tools, certifications, metrics, industries, scope, or outcomes.
4. Prefer exact selected keyword phrases where natural.
5. If none of the selected keywords fit, still return the best truthful rewrite.
6. Avoid keyword stuffing and AI-sounding phrasing.
7. Preserve the original meaning and resume tone.
8. You may rewrite the FULL line if needed.
9. Do NOT exceed {max_chars if max_chars is not None else "the original"} characters.
10. Stay as close as possible to the original line length.
11. The output must remain layout-safe for a tightly formatted resume.
12. {"Preserve the FIRST WORD exactly: " + first_word if preserve_first_word else "Keep the opening style natural and consistent with the original."}
13. Prefer replacing weaker words with stronger ATS language instead of adding extra words.
14. Return 1 to 5 options. More good options are preferred, but do not force weak ones.

Return ONLY valid JSON:
{{
  "line_index": {line_index},
  "original": {json.dumps(original_line, ensure_ascii=False)},
  "analysis": {{
    "selected_keywords_considered": {json.dumps(selected_keywords, ensure_ascii=False)},
    "selected_keywords_used": [],
    "selected_keywords_not_used": [],
    "main_gap": "",
    "risk_of_overclaiming": []
  }},
  "options": [
{options_schema}
  ]
}}

Context before: {json.dumps(surrounding_before, ensure_ascii=False)}
Target line: {json.dumps(original_line, ensure_ascii=False)}
Context after: {json.dumps(surrounding_after, ensure_ascii=False)}

SELECTED KEYWORDS:
{json.dumps(selected_keywords, ensure_ascii=False)}

Job description:
{job_description}
""".strip()

    @staticmethod
    def build_best_option_selector_prompt(
        original_line: str,
        options: List[str],
        max_chars: Optional[int] = None,
    ) -> str:
        options_block = "\n".join([f"{i + 1}. {opt}" for i, opt in enumerate(options)])
        char_rule = (
            f"Do not exceed {max_chars} characters."
            if max_chars is not None
            else "Keep the length close to the original."
        )

        return f"""
Choose the best rewritten resume line.

Return ONLY valid JSON:
{{
  "best_option_number": 1,
  "reason": "",
  "rejected_option_notes": {{}}
}}

Original:
{original_line}

Options:
{options_block}

Constraint:
{char_rule}
""".strip()


class ATSReviewEngine:
    """
    Orchestrates:
    1. JD deep parse
    2. Resume line relevance scoring
    3. Line rewrite generation
    4. Best-option selection
    5. Global gap + parser-risk audit
    """

    def __init__(self, processor: ResumeProcessor):
        self.processor = processor

    def build_jd_parse_prompt(self, job_description: str) -> str:
        return ResumeRewritePrompts.build_jd_deep_parse_prompt(job_description)

    def parse_jd_analysis(self, raw_model_output: str) -> JDAnalysis:
        data = ATSUtils.safe_parse_json(raw_model_output)
        return JDAnalysis(
            core_role_function=data.get("core_role_function", ""),
            primary_business_objective=data.get("primary_business_objective", ""),
            required_skills=ATSUtils.dedupe_keep_order(data.get("required_skills", [])),
            preferred_skills=ATSUtils.dedupe_keep_order(data.get("preferred_skills", [])),
            required_tools=ATSUtils.dedupe_keep_order(data.get("required_tools", [])),
            required_domain_terms=ATSUtils.dedupe_keep_order(data.get("required_domain_terms", [])),
            required_responsibilities=ATSUtils.dedupe_keep_order(data.get("required_responsibilities", [])),
            required_qualifications=ATSUtils.dedupe_keep_order(data.get("required_qualifications", [])),
            preferred_qualifications=ATSUtils.dedupe_keep_order(data.get("preferred_qualifications", [])),
            repeated_keywords=ATSUtils.dedupe_keep_order(data.get("repeated_keywords", [])),
            synonyms_and_variants=data.get("synonyms_and_variants", {}) or {},
            implied_screening_signals=ATSUtils.dedupe_keep_order(data.get("implied_screening_signals", [])),
            metrics_kpi_language=ATSUtils.dedupe_keep_order(data.get("metrics_kpi_language", [])),
            ats_critical_terms=ATSUtils.dedupe_keep_order(data.get("ats_critical_terms", [])),
            required_keywords=ATSUtils.dedupe_keep_order(data.get("required_keywords", [])),
            preferred_keywords=ATSUtils.dedupe_keep_order(data.get("preferred_keywords", [])),
            high_priority_missing=ATSUtils.dedupe_keep_order(data.get("high_priority_missing", [])),
            medium_priority_missing=ATSUtils.dedupe_keep_order(data.get("medium_priority_missing", [])),
            low_priority_missing=ATSUtils.dedupe_keep_order(data.get("low_priority_missing", [])),
            recommended_keyword_targets=ATSUtils.dedupe_keep_order(data.get("recommended_keyword_targets", [])),
        )

    def get_relevant_lines(self, jd: JDAnalysis, include_low: bool = False) -> List[LineSuggestion]:
        suggestions: List[LineSuggestion] = []
        for line in self.processor.get_all_lines(include_empty=False):
            assessment = ATSScorer.assess_line_relevance(line["text"], jd)
            if include_low or assessment.relevance in {"medium", "high"}:
                suggestions.append(
                    LineSuggestion(
                        line_index=line["index"],
                        original=line["text"],
                        assessment=assessment,
                    )
                )
        return suggestions

    def build_line_prompt(
        self,
        line_index: int,
        job_description: str,
        jd: JDAnalysis,
        selected_keywords: Optional[List[str]] = None,
    ) -> str:
        line = self.processor.get_line(line_index)
        if line is None:
            raise IndexError(f"Invalid line_index: {line_index}")

        context = self.processor.get_context_window(line_index, window=1)
        max_chars = line.char_count

        chosen_keywords = selected_keywords or jd.recommended_keyword_targets or jd.high_priority_missing or []

        return ResumeRewritePrompts.build_line_rewrite_prompt(
            line_index=line_index,
            original_line=line.text,
            job_description=job_description,
            selected_keywords=chosen_keywords[:12],
            surrounding_before=context.get("before", []),
            surrounding_after=context.get("after", []),
            max_chars=max_chars,
            n_options=5,
        )

    def parse_line_options(
        self,
        raw_model_output: str,
        jd: JDAnalysis,
        original_line: str,
        max_chars: Optional[int] = None,
        selected_keywords: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        data = ATSUtils.safe_parse_json(raw_model_output)
        options: List[RewriteOption] = []

        raw_options = data.get("options", [])
        if not isinstance(raw_options, list):
            raw_options = []

        original_normalized = ATSUtils.normalize_text(original_line)
        original_first = ATSUtils.first_word(original_line)
        preserve_first = ATSUtils.should_preserve_first_word(original_line)
        min_chars = max(1, len(original_line) - 8)

        seen = set()

        for option in raw_options:
            if not isinstance(option, dict):
                continue

            text = ATSUtils.clean_text(option.get("text", ""))
            if not text:
                continue

            text_norm = ATSUtils.normalize_text(text)
            if not text_norm:
                continue

            if text_norm.lower() == original_normalized.lower():
                continue

            if text_norm.lower() in seen:
                continue
            seen.add(text_norm.lower())

            if max_chars is not None and len(text) > max_chars:
                continue

            if len(text) < min_chars:
                continue

            if preserve_first:
                option_first = ATSUtils.first_word(text)
                if option_first != original_first:
                    continue

            keyword_hits = ATSUtils.find_keyword_hits(text, selected_keywords or [])
            keywords_added = ATSUtils.dedupe_keep_order(option.get("keywords_added", []))
            keywords_added = [kw for kw in keywords_added if ATSUtils.normalize_token(kw) in {
                ATSUtils.normalize_token(k) for k in (selected_keywords or [])
            }]

            if not keywords_added:
                keywords_added = keyword_hits

            score = ATSScorer.score_rewrite_option(
                original=original_line,
                option=text,
                jd=jd,
                max_chars=max_chars,
                selected_keywords=selected_keywords or [],
            )

            options.append(
                RewriteOption(
                    text=text,
                    keywords_added=ATSUtils.dedupe_keep_order(keywords_added),
                    score=score,
                    why_it_works=option.get("why_it_works", ""),
                )
            )

        options = sorted(options, key=lambda x: x.score, reverse=True)[:5]

        return {
            "line_index": data.get("line_index"),
            "original": data.get("original", original_line),
            "analysis": data.get("analysis", {}),
            "options": options,
        }

    def choose_best_option(self, original_line: str, options: List[RewriteOption]) -> Dict[str, Any]:
        if not options:
            return {
                "best_option_number": 1,
                "reason": "No valid options returned.",
                "rejected_option_notes": {},
            }

        ranked = sorted(enumerate(options, start=1), key=lambda x: x[1].score, reverse=True)
        best_num, best_obj = ranked[0]

        rejected = {}
        for idx, obj in ranked[1:]:
            rejected[str(idx)] = f"Lower score ({obj.score}) than best option ({best_obj.score})."

        return {
            "best_option_number": best_num,
            "reason": f"Best blend of ATS alignment, truthfulness, readability, and strict length safety with score {best_obj.score}.",
            "rejected_option_notes": rejected,
        }

    def build_master_prompt(self, job_description: str) -> str:
        return ResumeRewritePrompts.build_master_ats_prompt(
            resume_lines=self.processor.get_all_lines(include_empty=False),
            job_description=job_description,
        )

    def build_global_feedback(self, jd: JDAnalysis, reviewed_lines: List[LineSuggestion]) -> GlobalResumeFeedback:
        scores = [s.assessment.current_match_score for s in reviewed_lines]
        overall = int(sum(scores) / len(scores)) if scores else 0

        rewrite_based_improvements = [
            "Align the summary section with the exact job family and function.",
            "Mirror required tools and process language in the skills section where truthful.",
            "Strengthen bullets using action + method + result structure within the same line length.",
            "Add exact JD terminology only where naturally supported by the experience and layout permits.",
        ]

        substantive_gaps = []
        parser_risks = ATSUtils.COMMON_PARSER_RISKS.copy()

        if jd.required_tools:
            substantive_gaps.append(
                f"Verify actual experience with required tools: {', '.join(jd.required_tools[:8])}."
            )
        if jd.required_qualifications:
            substantive_gaps.append(
                f"Check coverage of required qualifications: {', '.join(jd.required_qualifications[:6])}."
            )

        blockers = []
        if jd.required_qualifications:
            blockers.append("missing required qualifications cannot be solved by rewriting alone")
        if jd.required_tools:
            blockers.append("missing required tools or systems cannot be claimed without real experience")
        blockers.append("strict line-length preservation limits how many keywords can be added")
        blockers.append("ATS match varies by employer system and weighting logic")

        return GlobalResumeFeedback(
            rewrite_based_improvements=ATSUtils.dedupe_keep_order(rewrite_based_improvements),
            substantive_gaps=ATSUtils.dedupe_keep_order(substantive_gaps),
            parser_risks=ATSUtils.dedupe_keep_order(parser_risks),
            overall_estimated_match_score=overall,
            what_prevents_true_100_percent_match=ATSUtils.dedupe_keep_order(blockers),
        )

    def export_review_bundle(self, bundle: ATSReviewBundle) -> Dict[str, Any]:
        return {
            "jd_analysis": asdict(bundle.jd_analysis),
            "line_suggestions": [asdict(item) for item in bundle.line_suggestions],
            "global_resume_feedback": asdict(bundle.global_resume_feedback),
        }


# =========================
# Example usage
# =========================

EXAMPLE_USAGE = r"""
with open("resume.docx", "rb") as f:
    processor = ResumeProcessor(f.read())

engine = ATSReviewEngine(processor)

jd_prompt = engine.build_jd_parse_prompt(job_description)
jd = engine.parse_jd_analysis(jd_model_output)
relevant_lines = engine.get_relevant_lines(jd)

for suggestion in relevant_lines:
    prompt = engine.build_line_prompt(
        suggestion.line_index,
        job_description,
        jd,
        selected_keywords=jd.recommended_keyword_targets[:12]
    )

    parsed = engine.parse_line_options(
        raw_model_output=line_model_output,
        jd=jd,
        original_line=suggestion.original,
        max_chars=len(suggestion.original),
        selected_keywords=jd.recommended_keyword_targets[:12],
    )

    suggestion.options = parsed["options"]

    best = engine.choose_best_option(suggestion.original, suggestion.options)
    suggestion.best_option_number = best["best_option_number"]
    suggestion.best_option_reason = best["reason"]

feedback = engine.build_global_feedback(jd, relevant_lines)
bundle = ATSReviewBundle(
    jd_analysis=jd,
    line_suggestions=relevant_lines,
    global_resume_feedback=feedback,
)

result_json = engine.export_review_bundle(bundle)
print(json.dumps(result_json, indent=2))
"""

if __name__ == "__main__":
    print("Enhanced ATS resume processor loaded.")
