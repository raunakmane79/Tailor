from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class ResumeLine:
    index: int
    text: str
    char_count: int
    source: str  # "paragraph" | "table"
    para_index: int = -1
    table_ref: Optional[Tuple[int, int, int]] = None
    is_empty: bool = False


class ResumeProcessor:
    """
    Parses a DOCX resume into line-level units and supports precise replacement
    while preserving formatting as much as possible.
    """

    def __init__(self, docx_bytes: bytes):
        self._original_bytes = docx_bytes
        self.doc = Document(io.BytesIO(docx_bytes))
        self._lines: List[ResumeLine] = []
        self._line_lookup: Dict[int, Dict[str, Any]] = {}
        self._parse()

    @staticmethod
    def _clean_text(text: str) -> str:
        if not text:
            return ""
        text = text.replace("\xa0", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n+", "\n", text)
        return text.strip()

    def _add_line(self, line: ResumeLine) -> None:
        self._lines.append(line)
        lookup: Dict[str, Any] = {"source": line.source}

        if line.source == "paragraph":
            lookup["para_index"] = line.para_index
        elif line.source == "table":
            lookup["table_ref"] = line.table_ref

        self._line_lookup[line.index] = lookup

    def _parse(self) -> None:
        self._lines = []
        self._line_lookup = {}
        idx = 0

        for pi, para in enumerate(self.doc.paragraphs):
            text = self._clean_text(para.text)
            self._add_line(
                ResumeLine(
                    index=idx,
                    text=text,
                    char_count=len(text),
                    source="paragraph",
                    para_index=pi,
                    is_empty=(text == ""),
                )
            )
            idx += 1

        for ti, table in enumerate(self.doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    # Join all paragraphs in a cell so extracted text reflects
                    # what the user actually sees in the cell.
                    cell_text = "\n".join(
                        self._clean_text(p.text) for p in cell.paragraphs if self._clean_text(p.text)
                    )
                    cell_text = self._clean_text(cell_text)

                    self._add_line(
                        ResumeLine(
                            index=idx,
                            text=cell_text,
                            char_count=len(cell_text),
                            source="table",
                            table_ref=(ti, ri, ci),
                            is_empty=(cell_text == ""),
                        )
                    )
                    idx += 1

    def get_all_lines(self, include_empty: bool = True) -> List[Dict[str, Any]]:
        lines = self._lines if include_empty else [l for l in self._lines if not l.is_empty]
        return [
            {
                "index": l.index,
                "text": l.text,
                "char_count": l.char_count,
                "source": l.source,
            }
            for l in lines
        ]

    def get_line(self, line_index: int) -> Optional[ResumeLine]:
        if 0 <= line_index < len(self._lines):
            return self._lines[line_index]
        return None

    def get_context_window(self, line_index: int, window: int = 1) -> Dict[str, Any]:
        """
        Returns nearby lines so the LLM can match tone with minimal token use.
        """
        if not (0 <= line_index < len(self._lines)):
            return {"target": "", "before": [], "after": []}

        start = max(0, line_index - window)
        end = min(len(self._lines), line_index + window + 1)

        return {
            "target": self._lines[line_index].text,
            "before": [self._lines[i].text for i in range(start, line_index)],
            "after": [self._lines[i].text for i in range(line_index + 1, end)],
        }

    def replace_line(self, line_index: int, new_text: str) -> bool:
        if not (0 <= line_index < len(self._lines)):
            return False

        new_text = self._clean_text(new_text)
        line = self._lines[line_index]
        lookup = self._line_lookup[line_index]

        if lookup["source"] == "paragraph":
            para = self.doc.paragraphs[lookup["para_index"]]
            self._replace_para_text(para, new_text)

        elif lookup["source"] == "table":
            ti, ri, ci = lookup["table_ref"]
            cell = self.doc.tables[ti].rows[ri].cells[ci]
            self._replace_cell_text(cell, new_text)

        line.text = new_text
        line.char_count = len(new_text)
        line.is_empty = (new_text == "")
        return True

    def _replace_cell_text(self, cell, new_text: str) -> None:
        """
        Replace all visible text in a table cell while preserving style from the
        first meaningful paragraph/run when possible.
        """
        if cell.paragraphs:
            target_para = cell.paragraphs[0]
            self._replace_para_text(target_para, new_text)

            # Clear any remaining paragraphs in the cell so old text does not linger.
            for extra_para in cell.paragraphs[1:]:
                self._replace_para_text(extra_para, "")
        else:
            cell.text = new_text

    @staticmethod
    def _replace_para_text(para, new_text: str) -> None:
        """
        Replace paragraph text while preserving the first meaningful run style.
        """
        template_run = None
        for run in para.runs:
            if run.text and run.text.strip():
                template_run = run
                break

        if template_run is None and para.runs:
            template_run = para.runs[0]

        rpr_clone = None
        if template_run is not None:
            rpr = template_run._r.find(qn("w:rPr"))
            if rpr is not None:
                rpr_clone = copy.deepcopy(rpr)

        p_elem = para._p

        # Remove all existing runs
        for child in list(p_elem):
            if child.tag == qn("w:r"):
                p_elem.remove(child)

        new_run = OxmlElement("w:r")
        if rpr_clone is not None:
            new_run.append(rpr_clone)

        text_elem = OxmlElement("w:t")
        if new_text.startswith(" ") or new_text.endswith(" "):
            text_elem.set(qn("xml:space"), "preserve")
        text_elem.text = new_text

        new_run.append(text_elem)
        p_elem.append(new_run)

    def export(self) -> bytes:
        output = io.BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()


class ResumeRewritePrompts:
    """
    Builds compact prompts for natural, high-quality rewriting with low token use.
    """

    @staticmethod
    def build_line_rewrite_prompt(
        original_line: str,
        job_description: str,
        surrounding_before: Optional[List[str]] = None,
        surrounding_after: Optional[List[str]] = None,
        max_chars: Optional[int] = None,
        keywords_to_include: Optional[List[str]] = None,
        n_options: int = 3,
    ) -> str:
        surrounding_before = surrounding_before or []
        surrounding_after = surrounding_after or []
        keywords_to_include = keywords_to_include or []

        context_before = " | ".join([x for x in surrounding_before if x][:1])
        context_after = " | ".join([x for x in surrounding_after if x][:1])
        keywords_str = ", ".join(keywords_to_include[:8])

        char_rule = (
            f"Keep length <= {max_chars} chars."
            if max_chars is not None
            else "Keep length close to original."
        )
        keyword_rule = (
            f"Use these keywords only if natural: {keywords_str}."
            if keywords_str
            else "Add job-description keywords only if natural."
        )

        return f"""
Rewrite this resume line for better ATS match.

Rules:
- Keep meaning truthful.
- Match the tone and style of the original sentence.
- Sound natural, not AI-written.
- Be concise and specific.
- {char_rule}
- Do not use filler or vague buzzwords.
- Preserve resume style.
- {keyword_rule}
- Return ONLY {n_options} rewritten options as a JSON array of strings.

Context before: {context_before}
Target line: {original_line}
Context after: {context_after}

Job description:
{job_description}
""".strip()

    @staticmethod
    def build_keyword_extraction_prompt(job_description: str, max_keywords: int = 15) -> str:
        return f"""
Extract the most important ATS keywords from this job description.

Rules:
- Prefer hard skills, tools, domain phrases, methods, and business terms.
- Avoid generic soft skills unless repeated or clearly important.
- Merge duplicates.
- Return ONLY valid JSON in this format:
{{"keywords":["kw1","kw2","kw3"]}}
- Max {max_keywords} keywords.

Job description:
{job_description}
""".strip()

    @staticmethod
    def build_best_option_selector_prompt(
        original_line: str,
        options: List[str],
        max_chars: Optional[int] = None,
    ) -> str:
        options_block = "\n".join(f"{i + 1}. {opt}" for i, opt in enumerate(options))
        char_rule = (
            f"Must be <= {max_chars} chars."
            if max_chars is not None
            else "Keep length close to original."
        )

        return f"""
Pick the best rewritten resume line.

Rules:
- Must sound the most natural.
- Must preserve the original meaning.
- Must fit resume style.
- Should be stronger and clearer than the original.
- {char_rule}
- Return ONLY valid JSON:
{{"best_option_number": 1, "reason": "short reason"}}

Original:
{original_line}

Options:
{options_block}
""".strip()
